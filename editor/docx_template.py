from __future__ import annotations

import io
import re
from copy import deepcopy
from pathlib import Path
from typing import Any, TypeAlias

from docx import Document as DocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.run import Run

DocxDocumentType: TypeAlias = Any
ContactPart: TypeAlias = tuple[str, str, str | None, int]

# Meses para formatear fechas (YYYY-MM -> Mes YYYY)
MONTHS_REVERSE = {
    "01": "Ene",
    "02": "Feb",
    "03": "Mar",
    "04": "Abr",
    "05": "May",
    "06": "Jun",
    "07": "Jul",
    "08": "Ago",
    "09": "Sep",
    "10": "Oct",
    "11": "Nov",
    "12": "Dic",
}


def _format_detail_line(value: str | None) -> str:
    detail = (value or "").strip()
    return detail


def _normalize_extra_mode(value: str | None, default: str = "subtitles") -> str:
    """Normaliza modos de extras.

    - Modo histórico "items" se trata como alias de "subtitles" (Punto/Listado).
    - Modos válidos: subtitles, subtitle_items, detailed.
    """
    mode = (value or "").strip() or default
    if mode == "items":
        return "subtitles"
    if mode not in {"subtitles", "subtitle_items", "detailed"}:
        return default
    return mode


def _entry_items_inline(entry: dict) -> str:
    items = [str(item).strip() for item in (entry.get("items") or []) if str(item).strip()]
    return ", ".join(items)


def render_from_template(structured: dict, template_path: Path, font_name: str | None = None) -> bytes:
    # Carga la plantilla DOCX y reemplaza secciones con la data estructurada
    doc = DocxDocument(str(template_path))
    if not doc.tables:
        raise ValueError("La plantilla no contiene tablas.")

    table = doc.tables[0]
    basics = structured.get("basics", {})
    experience = structured.get("experience", []) or []
    education = structured.get("education", []) or []
    skills = structured.get("skills", []) or []
    extras = structured.get("extra_sections", []) or []

    exp_header_idx = _find_row_index(table, "experiencia")
    edu_header_idx = _find_row_index(table, "educación")
    skills_header_idx = _find_row_index(table, "habilidades")
    if exp_header_idx is None or edu_header_idx is None or skills_header_idx is None:
        raise ValueError("No se encontraron secciones clave en la plantilla.")

    contact_row_idx = _find_row_index_predicate(
        table,
        lambda text: "linkedin" in text or "github" in text or "@" in text,
    )
    name_row_idx = _find_first_non_empty_before(table, contact_row_idx or exp_header_idx)
    summary_row_idx = _find_next_non_empty_row(table, (contact_row_idx or 0) + 1, exp_header_idx)

    if name_row_idx is not None:
        _set_row_text(table.rows[name_row_idx], basics.get("name", "").strip())

    if contact_row_idx is not None:
        _set_contact_row(table.rows[contact_row_idx], basics)

    if summary_row_idx is not None:
        _set_row_text(table.rows[summary_row_idx], basics.get("description", "").strip())

    # Experiencia y educacion se escriben en bloques con filas clonadas
    _set_row_keep_with_next(table.rows[exp_header_idx])
    _set_row_keep_with_next(table.rows[edu_header_idx])
    _apply_section_keep_with_next_gap(table, exp_header_idx)
    _apply_experience(table, exp_header_idx, edu_header_idx, experience)

    edu_header_idx = _find_row_index(table, "educación")
    skills_header_idx = _find_row_index(table, "habilidades")
    if edu_header_idx is None or skills_header_idx is None:
        raise ValueError("No se encontraron secciones clave en la plantilla.")

    _set_row_keep_with_next(table.rows[edu_header_idx])
    _set_row_keep_with_next(table.rows[skills_header_idx])
    _apply_section_keep_with_next_gap(table, edu_header_idx)
    _apply_education(table, edu_header_idx, skills_header_idx, education)

    skills_header_idx = _find_row_index(table, "habilidades")
    if skills_header_idx is None:
        raise ValueError("No se encontro la seccion de habilidades en la plantilla.")

    _set_row_keep_with_next(table.rows[skills_header_idx])
    _apply_section_keep_with_next_gap(table, skills_header_idx)
    skills_content_idx = _apply_skills(table, skills_header_idx, skills)

    # Extras se agregan despues de habilidades (ahora soporta modo por entrada)
    extra_blocks = _apply_extras(table, skills_header_idx, skills_content_idx, extras)

    # Fuerza tamaño de bullets en habilidades
    _normalize_skills_bullets(doc, table, skills_content_idx, size_pt=11)
    # Reordena módulos según core_order enviado por la UI (flechas)
    _apply_module_order(table, structured, extra_blocks)
    # Aplica fuente global si el usuario la selecciono
    _apply_font(doc, font_name)
    _collapse_blank_rows(table)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _apply_experience(table, exp_header_idx: int, edu_header_idx: int, experience: list[dict]) -> None:
    # Inserta bloques de experiencia usando filas template
    exp_role_idx = _find_next_non_empty_row(table, exp_header_idx + 1, edu_header_idx)
    exp_high_idx = _find_next_non_empty_row(table, (exp_role_idx or exp_header_idx) + 1, edu_header_idx)
    if exp_role_idx is None or exp_high_idx is None:
        return

    role_template = deepcopy(table.rows[exp_role_idx]._tr)
    high_template = deepcopy(table.rows[exp_high_idx]._tr)
    spacer_idx = _find_blank_row_without_borders(table, exp_high_idx + 1, edu_header_idx)
    if spacer_idx is None:
        spacer_idx = _find_any_blank_row_without_borders(table)
    section_spacer_idx = _find_trailing_blank_row_without_borders(table, exp_high_idx + 1, edu_header_idx)
    if section_spacer_idx is None:
        section_spacer_idx = _find_any_blank_row_without_borders(table)
    spacer_template = deepcopy(table.rows[spacer_idx]._tr) if spacer_idx is not None else None
    section_spacer_template = (
        deepcopy(table.rows[section_spacer_idx]._tr) if section_spacer_idx is not None else None
    )

    gap_idx = exp_header_idx + 1
    has_gap = gap_idx < edu_header_idx and _is_blank_row(table.rows[gap_idx])
    if has_gap:
        _set_row_keep_with_next(table.rows[gap_idx])
        _clear_row_height(table.rows[gap_idx])
    items = [item for item in (experience or []) if _has_experience_content(item)]
    delete_start = gap_idx + 1 if has_gap and items else gap_idx
    _remove_rows(table, delete_start, edu_header_idx)

    if not items:
        return

    insert_idx = delete_start
    for idx, item in enumerate(items):
        _insert_row_before(table, insert_idx, deepcopy(role_template))
        role_row = table.rows[insert_idx]
        _clear_row_height(role_row)
        _fill_experience_role_row(role_row, item)
        insert_idx += 1

        has_highlights = any((h or "").strip() for h in (item.get("highlights") or []))
        if has_highlights:
            _set_row_keep_with_next(role_row)
            _insert_row_before(table, insert_idx, deepcopy(high_template))
            high_row = table.rows[insert_idx]
            _clear_row_height(high_row)
            _fill_experience_highlights_row(high_row, item)
            insert_idx += 1

        if spacer_template is not None and idx < len(items) - 1:
            _insert_row_before(table, insert_idx, deepcopy(spacer_template))
            _clear_row_height(table.rows[insert_idx])
            insert_idx += 1
    if section_spacer_template is not None and items:
        _insert_row_before(table, insert_idx, deepcopy(section_spacer_template))
        _clear_row_height(table.rows[insert_idx])
        insert_idx += 1


def _apply_education(table, edu_header_idx: int, skills_header_idx: int, education: list[dict]) -> None:
    # Inserta bloques de educacion usando filas template
    edu_template_idx = _find_next_non_empty_row(table, edu_header_idx + 1, skills_header_idx)
    if edu_template_idx is None:
        return

    template_tr = deepcopy(table.rows[edu_template_idx]._tr)
    spacer_idx = _find_blank_row_without_borders(table, edu_template_idx + 1, skills_header_idx)
    if spacer_idx is None:
        spacer_idx = _find_any_blank_row_without_borders(table)
    section_spacer_idx = _find_trailing_blank_row_without_borders(table, edu_template_idx + 1, skills_header_idx)
    if section_spacer_idx is None:
        section_spacer_idx = _find_any_blank_row_without_borders(table)
    spacer_template = deepcopy(table.rows[spacer_idx]._tr) if spacer_idx is not None else None
    section_spacer_template = (
        deepcopy(table.rows[section_spacer_idx]._tr) if section_spacer_idx is not None else None
    )

    gap_idx = edu_header_idx + 1
    has_gap = gap_idx < skills_header_idx and _is_blank_row(table.rows[gap_idx])
    if has_gap:
        _set_row_keep_with_next(table.rows[gap_idx])
        _clear_row_height(table.rows[gap_idx])
    items = [item for item in (education or []) if _has_education_content(item)]
    delete_start = gap_idx + 1 if has_gap and items else gap_idx
    _remove_rows(table, delete_start, skills_header_idx)

    if not items:
        return

    insert_idx = delete_start
    for idx, item in enumerate(items):
        _insert_row_before(table, insert_idx, deepcopy(template_tr))
        row = table.rows[insert_idx]
        _clear_row_height(row)
        _fill_education_row(row, item)
        insert_idx += 1
        if spacer_template is not None and idx < len(items) - 1:
            _insert_row_before(table, insert_idx, deepcopy(spacer_template))
            _clear_row_height(table.rows[insert_idx])
            insert_idx += 1
    if section_spacer_template is not None and items:
        _insert_row_before(table, insert_idx, deepcopy(section_spacer_template))
        _clear_row_height(table.rows[insert_idx])
        insert_idx += 1


def _apply_skills(table, skills_header_idx: int, skills: list[dict]) -> int | None:
    # Rellena la celda de habilidades respetando estilos
    skills_content_idx = _find_next_non_empty_row(table, skills_header_idx + 1, None)
    if skills_content_idx is None:
        return None

    cell = _unique_cells(table.rows[skills_content_idx])[0]
    template_paragraphs = list(cell.paragraphs)
    template_paragraph = template_paragraphs[0] if template_paragraphs else None
    template_style = template_paragraph.style if template_paragraph is not None else None
    _clear_cell(cell)

    filtered = [item for item in skills if (item.get("category") or item.get("items"))]
    if not filtered:
        _add_paragraph(cell, "", style=template_style, paragraph_template=template_paragraph)
        return skills_content_idx

    category_template = None
    items_template = None
    spacer_template = None
    for para in template_paragraphs:
        if para.text.strip():
            if category_template is None:
                category_template = para
            elif items_template is None:
                items_template = para
        elif spacer_template is None and category_template is not None:
            spacer_template = para
        if category_template and items_template and spacer_template:
            break
    if category_template is None:
        category_template = template_paragraph
    if items_template is None:
        items_template = category_template
    if spacer_template is None:
        spacer_template = items_template

    for idx, item in enumerate(filtered):
        category = (item.get("category") or "").strip()
        items = (item.get("items") or "").strip()
        if category:
            _add_paragraph(
                cell,
                category,
                style=template_style,
                run_template=category_template.runs[0] if category_template is not None and category_template.runs else None,
                paragraph_template=category_template,
            )
        if items:
            _add_paragraph(
                cell,
                items,
                style=template_style,
                run_template=items_template.runs[0] if items_template is not None and items_template.runs else None,
                paragraph_template=items_template,
            )
        if idx < len(filtered) - 1:
            _add_paragraph(
                cell,
                "",
                style=template_style,
                run_template=spacer_template.runs[0] if spacer_template is not None and spacer_template.runs else None,
                paragraph_template=spacer_template,
            )

    return skills_content_idx


def _normalize_skills_bullets(
    doc: DocxDocumentType,
    table,
    skills_content_idx: int | None,
    *,
    size_pt: int = 11,
) -> None:
    # Ajusta tamaño de la lista numerada/bullets en habilidades
    if skills_content_idx is None:
        return
    row = table.rows[skills_content_idx]
    num_pairs: set[tuple[str, str]] = set()
    for cell in _unique_cells(row):
        for paragraph in cell.paragraphs:
            ppr = paragraph._p.pPr
            if ppr is None or ppr.numPr is None:
                continue
            num_id = ppr.numPr.numId.val if ppr.numPr.numId is not None else None
            ilvl = ppr.numPr.ilvl.val if ppr.numPr.ilvl is not None else None
            if num_id is None or ilvl is None:
                continue
            num_pairs.add((str(num_id), str(ilvl)))

    for num_id, ilvl in num_pairs:
        _set_numbering_level_size(doc, num_id, ilvl, size_pt)


def _extra_entry_lines(entry: dict, *, mode: str = "detailed") -> list[str]:
    """Devuelve líneas de una entry para compatibilidad.

    Nota: el render real en DOCX (extras) se hace en _apply_extras() y soporta
    modo por entrada. Esta función queda como helper de fallback.
    """
    normalized_mode = _normalize_extra_mode(mode, default="detailed")
    subtitle = (entry.get("subtitle") or "").strip()
    items_inline = _entry_items_inline(entry)

    # Modo "skills-like"
    # "subtitles" ahora representa "Punto/Listado".
    # Compatibilidad: si viene data antigua con items pero sin subtitle, usamos items_inline.
    if normalized_mode == "subtitles":
        text = subtitle or items_inline
        return [text] if text else []
    if normalized_mode == "subtitle_items":
        lines: list[str] = []
        if subtitle:
            lines.append(subtitle)
        if items_inline:
            lines.append(items_inline)
        return lines

    # detailed (comportamiento histórico)
    title = (entry.get("title") or "").strip()
    where = (entry.get("where") or "").strip()
    header = " — ".join([part for part in [title, where] if part])
    date_range = _format_date_range(entry.get("start", ""), entry.get("end", ""))
    location = _join_location(entry.get("city"), entry.get("country"))
    tech = (entry.get("tech") or "").strip()
    tech_line = _format_detail_line(tech)
    lines: list[str] = []
    if subtitle:
        lines.append(subtitle)
    if header:
        lines.append(f"{header} ({date_range})" if date_range else header)
    elif date_range:
        lines.append(date_range)
    if tech_line:
        lines.append(tech_line)
    if location:
        lines.append(location)
    items = [item for item in (entry.get("items") or []) if item and str(item).strip()]
    lines.extend(items)
    return lines


def _extra_entry_has_content(entry: dict) -> bool:
    if not entry:
        return False
    fields = [
        entry.get("subtitle"),
        entry.get("title"),
        entry.get("where"),
        entry.get("tech"),
        entry.get("start"),
        entry.get("end"),
        entry.get("city"),
        entry.get("country"),
    ]
    if any(field and str(field).strip() for field in fields):
        return True
    return any(item for item in (entry.get("items") or []) if str(item).strip())


def _flatten_extra_lines(extra: dict) -> list[str]:
    """Convierte una extra-section a lista de líneas (fallback/compatibilidad).

    - Si viene con el formato antiguo (title/items) devuelve items.
    - Si viene con entries, usa mode por entrada (o el mode de la sección como default).
    """
    if "entries" not in extra:
        return [item for item in (extra.get("items") or []) if item and str(item).strip()]

    default_mode = _normalize_extra_mode(extra.get("mode"), default="subtitles")
    lines: list[str] = []
    for entry in extra.get("entries") or []:
        entry_mode = _normalize_extra_mode(entry.get("mode"), default=default_mode)
        lines.extend(_extra_entry_lines(entry, mode=entry_mode))
    return [line for line in lines if line and str(line).strip()]


def _apply_extras(
    table,
    skills_header_idx: int,
    skills_content_idx: int | None,
    extras: list[dict],
) -> list[dict[str, Any]]:
    # Crea nuevas secciones extras bajo habilidades.
    #
    # Soporta modo por ENTRADA (entry["mode"]) para permitir mezclar dentro de una misma sección.
    # Compatibilidad:
    # - si no hay entry.mode, cae en extra.mode o "items".
    if skills_content_idx is None or not extras:
        return []

    filtered: list[dict] = []
    for extra_idx, extra in enumerate(extras, start=1):
        section_id = (extra.get("section_id") or "").strip()
        raw_title = (extra.get("title") or "").strip()
        # Si el usuario deja el título vacío, forzamos un fallback para
        # preservar límites de sección al reimportar desde PDF.
        title = raw_title or f"SECCION EXTRA {extra_idx}"
        section_mode = (extra.get("mode") or "subtitles").strip() or "subtitles"
        section_mode = _normalize_extra_mode(section_mode, default="subtitles")
        entries = [entry for entry in (extra.get("entries") or []) if _extra_entry_has_content(entry)]
        # Normaliza entry.mode (fallback al modo de sección)
        for entry in entries:
            entry_mode = (entry.get("mode") or section_mode).strip() or section_mode
            entry["mode"] = _normalize_extra_mode(entry_mode, default=section_mode)
        # Si viene en formato antiguo (sin entries), lo tratamos como lista de items
        if not entries and "items" in extra:
            items = [item for item in (extra.get("items") or []) if item and str(item).strip()]
            if title or items:
                filtered.append({"section_id": section_id, "title": title, "entries": [], "items": items})
            continue

        if title or entries:
            filtered.append({"section_id": section_id, "title": title, "entries": entries})

    if not filtered:
        return []

    header_template = deepcopy(table.rows[skills_header_idx]._tr)
    gap_idx = skills_header_idx + 1
    gap_template = None
    if gap_idx < len(table.rows) and _is_blank_row(table.rows[gap_idx]):
        gap_template = deepcopy(table.rows[gap_idx]._tr)

    # Template para contenido de 1 columna (como habilidades)
    content_template = deepcopy(table.rows[skills_content_idx]._tr)

    spacer_template = None
    spacer_idx = _find_any_blank_row_without_borders(table)
    if spacer_idx is not None:
        spacer_template = deepcopy(table.rows[spacer_idx]._tr)

    # Templates estilo "detalle" (2 columnas como experiencia)
    exp_role_template = None
    exp_high_template = None
    exp_header_idx = _find_row_index(table, "experiencia")
    edu_header_idx = _find_row_index(table, "educación")
    if exp_header_idx is not None:
        exp_role_idx = _find_next_non_empty_row(table, exp_header_idx + 1, edu_header_idx)
        exp_high_idx = _find_next_non_empty_row(table, (exp_role_idx or exp_header_idx) + 1, edu_header_idx)
        if exp_role_idx is not None:
            exp_role_template = deepcopy(table.rows[exp_role_idx]._tr)
        if exp_high_idx is not None:
            exp_high_template = deepcopy(table.rows[exp_high_idx]._tr)

    insert_idx = skills_content_idx + 1
    extra_blocks: list[dict[str, Any]] = []

    # Separación suave entre habilidades y la primera sección extra
    if spacer_template is not None:
        _insert_row_before(table, insert_idx, deepcopy(spacer_template))
        _clear_row_height(table.rows[insert_idx])
        insert_idx += 1

    for idx, extra in enumerate(filtered):
        section_id = (extra.get("section_id") or "").strip()
        title = extra.get("title", "")
        block_start = insert_idx

        _insert_row_before(table, insert_idx, deepcopy(header_template))
        header_row = table.rows[insert_idx]
        _clear_row_height(header_row)
        _set_row_text(header_row, title.upper() if title else "")
        _set_row_keep_with_next(header_row)
        insert_idx += 1

        if gap_template is not None:
            _insert_row_before(table, insert_idx, deepcopy(gap_template))
            gap_row = table.rows[insert_idx]
            _clear_row_height(gap_row)
            _set_row_keep_with_next(gap_row)
            insert_idx += 1

        entries = extra.get("entries") or []
        # Caso "antiguo": solo items (sin entries)
        if not entries and "items" in extra:
            _insert_row_before(table, insert_idx, deepcopy(content_template))
            content_row = table.rows[insert_idx]
            _clear_row_height(content_row)
            _fill_extra_row(content_row, [str(x).strip() for x in (extra.get("items") or []) if str(x).strip()])
            insert_idx += 1
        else:
            for entry_idx, entry in enumerate(entries):
                mode = _normalize_extra_mode((entry.get("mode") or "subtitles").strip() or "subtitles", default="subtitles")

                # detailed -> 2 columnas (si hay templates disponibles)
                if mode == "detailed" and exp_role_template is not None:
                    _insert_row_before(table, insert_idx, deepcopy(exp_role_template))
                    role_row = table.rows[insert_idx]
                    _clear_row_height(role_row)
                    _fill_extra_detail_role_row(role_row, entry)
                    insert_idx += 1

                    entry_items = [item for item in (entry.get("items") or []) if str(item).strip()]
                    if entry_items and exp_high_template is not None:
                        _set_row_keep_with_next(role_row)
                        _insert_row_before(table, insert_idx, deepcopy(exp_high_template))
                        high_row = table.rows[insert_idx]
                        _clear_row_height(high_row)
                        _fill_extra_detail_highlights_row(high_row, entry_items)
                        insert_idx += 1
                else:
                    # Cualquier otro modo -> 1 columna (como habilidades)
                    _insert_row_before(table, insert_idx, deepcopy(content_template))
                    content_row = table.rows[insert_idx]
                    _clear_row_height(content_row)
                    lines = _extra_entry_lines(entry, mode=mode)
                    _fill_extra_row(content_row, lines)
                    # Diferenciar visualmente subtítulos/puntos en Extras: en cursiva.
                    # Esto permite distinguir un "hito/subtítulo" de los items (bullets) al leer el PDF.
                    if mode in ("subtitles", "subtitle_items"):
                        cells = _unique_cells(content_row)
                        if cells and cells[0].paragraphs:
                            for run in cells[0].paragraphs[0].runs:
                                run.italic = True
                    insert_idx += 1

                if spacer_template is not None and entry_idx < len(entries) - 1:
                    _insert_row_before(table, insert_idx, deepcopy(spacer_template))
                    _clear_row_height(table.rows[insert_idx])
                    insert_idx += 1

        if spacer_template is not None and idx < len(filtered) - 1:
            _insert_row_before(table, insert_idx, deepcopy(spacer_template))
            _clear_row_height(table.rows[insert_idx])
            insert_idx += 1

        block_end = insert_idx
        block_rows = [table.rows[row_idx]._tr for row_idx in range(block_start, block_end)]
        if block_rows:
            extra_blocks.append({"section_id": section_id, "rows": block_rows})

    return extra_blocks


def _apply_module_order(table, structured: dict, extra_blocks: list[dict[str, Any]]) -> None:
    # Respeta el orden de módulos de la UI (core_order): experience, education, skills, extra-*
    meta = structured.get("meta") or {}
    raw_order = str(meta.get("core_order") or "").strip()
    if not raw_order:
        return

    requested_order = [item.strip() for item in raw_order.split(",") if item.strip()]
    if not requested_order:
        return

    exp_header_idx = _find_row_index(table, "experiencia")
    edu_header_idx = _find_row_index(table, "educación")
    if edu_header_idx is None:
        edu_header_idx = _find_row_index(table, "educacion")
    skills_header_idx = _find_row_index(table, "habilidades")
    if exp_header_idx is None or edu_header_idx is None or skills_header_idx is None:
        return

    tbl = table._tbl
    tr_list = list(tbl.tr_lst)
    index_by_tr_id = {id(tr): idx for idx, tr in enumerate(tr_list)}

    normalized_extra_blocks: list[dict[str, Any]] = []
    for idx, block in enumerate(extra_blocks or []):
        sid = (block.get("section_id") or "").strip() or f"extra-auto-{idx}"
        rows = [tr for tr in (block.get("rows") or []) if tr.getparent() is tbl]
        if not rows:
            continue
        positions = [index_by_tr_id[id(tr)] for tr in rows if id(tr) in index_by_tr_id]
        if not positions:
            continue
        normalized_extra_blocks.append({"section_id": sid, "rows": rows, "start": min(positions)})

    normalized_extra_blocks.sort(key=lambda block: block["start"])
    first_extra_start = normalized_extra_blocks[0]["start"] if normalized_extra_blocks else len(table.rows)

    exp_rows = [table.rows[idx]._tr for idx in range(exp_header_idx, min(edu_header_idx, len(table.rows)))]
    edu_rows = [table.rows[idx]._tr for idx in range(edu_header_idx, min(skills_header_idx, len(table.rows)))]
    skills_rows = [table.rows[idx]._tr for idx in range(skills_header_idx, min(first_extra_start, len(table.rows)))]

    module_rows: dict[str, list[Any]] = {
        "experience": exp_rows,
        "education": edu_rows,
        "skills": skills_rows,
    }

    extra_ids: list[str] = []
    for block in normalized_extra_blocks:
        sid = block["section_id"]
        module_rows[sid] = list(block["rows"])
        extra_ids.append(sid)

    current_modules = ["experience", "education", "skills"] + extra_ids

    expanded_requested: list[str] = []
    for module_id in requested_order:
        if module_id == "extras":
            expanded_requested.extend(extra_ids)
        else:
            expanded_requested.append(module_id)

    final_modules: list[str] = []
    seen_modules: set[str] = set()
    for module_id in expanded_requested:
        if module_id in module_rows and module_rows[module_id] and module_id not in seen_modules:
            final_modules.append(module_id)
            seen_modules.add(module_id)
    for module_id in current_modules:
        if module_id in module_rows and module_rows[module_id] and module_id not in seen_modules:
            final_modules.append(module_id)
            seen_modules.add(module_id)

    if not final_modules:
        return

    spacer_template = None
    spacer_idx = _find_any_blank_row_without_borders(table)
    if spacer_idx is not None:
        spacer_template = deepcopy(table.rows[spacer_idx]._tr)

    rows_to_move: list[Any] = []
    rows_to_move_ids: set[int] = set()
    for module_id in current_modules:
        for tr in module_rows.get(module_id, []):
            tr_id = id(tr)
            if tr_id in rows_to_move_ids:
                continue
            rows_to_move.append(tr)
            rows_to_move_ids.add(tr_id)

    if not rows_to_move:
        return

    insertion_positions = [idx for idx, tr in enumerate(tr_list) if id(tr) in rows_to_move_ids]
    if not insertion_positions:
        return
    insert_at = min(insertion_positions)

    for tr in tr_list:
        if id(tr) in rows_to_move_ids and tr.getparent() is tbl:
            tbl.remove(tr)

    inserted_rows: set[int] = set()
    cursor = insert_at
    for module_idx, module_id in enumerate(final_modules):
        module_inserted = False
        for tr in module_rows.get(module_id, []):
            tr_id = id(tr)
            if tr_id in inserted_rows:
                continue
            _insert_row_before(table, cursor, tr)
            cursor += 1
            inserted_rows.add(tr_id)
            module_inserted = True
        if (
            module_inserted
            and module_idx < len(final_modules) - 1
            and spacer_template is not None
            and cursor > 0
        ):
            previous_row = table.rows[cursor - 1]
            if not _is_blank_row(previous_row):
                _insert_row_before(table, cursor, deepcopy(spacer_template))
                _clear_row_height(table.rows[cursor])
                cursor += 1

def _fill_experience_role_row(row, item: dict) -> None:
    cells = _unique_cells(row)
    if not cells:
        return
    left = cells[0]
    right = cells[1] if len(cells) > 1 else None

    company = (item.get("company") or "").strip()
    role = (item.get("role") or "").strip()
    tech = (item.get("technologies") or "").strip()
    tech_line = _format_detail_line(tech)
    left_lines = _filter_empty_lines([company, role, tech_line])
    _set_cell_lines_preserve(left, left_lines, trim_extra_paragraphs=True)

    if right is not None:
        location = _join_location(item.get("city"), item.get("country"))
        date_range = _format_date_range(item.get("start"), item.get("end"))
        right_lines = _filter_empty_lines([location, date_range])
        _set_cell_lines_preserve(right, right_lines, trim_extra_paragraphs=True)


def _fill_extra_detail_role_row(row, entry: dict) -> None:
    cells = _unique_cells(row)
    if not cells:
        return
    left = cells[0]
    right = cells[1] if len(cells) > 1 else None

    title = (entry.get("title") or "").strip()
    where = (entry.get("where") or "").strip()
    tech = (entry.get("tech") or "").strip()
    tech_line = _format_detail_line(tech)

    # Orden similar a experiencia: "Empresa/Proyecto", "Rol/Curso", "Tecnologías/Detalle"
    left_lines: list[str] = []
    if where:
        left_lines.append(where)
    if title:
        left_lines.append(title)
    if tech_line:
        left_lines.append(tech_line)
    _set_cell_lines_preserve(left, _filter_empty_lines(left_lines), trim_extra_paragraphs=True)

    if right is not None:
        location = _join_location(entry.get("city"), entry.get("country"))
        date_range = _format_date_range(entry.get("start"), entry.get("end"))
        right_lines = _filter_empty_lines([location, date_range])
        _set_cell_lines_preserve(right, right_lines, trim_extra_paragraphs=True)


def _fill_extra_detail_highlights_row(row, items: list[str]) -> None:
    cells = _unique_cells(row)
    if not cells:
        return
    cell = cells[0]
    lines = _filter_empty_lines(items)
    _set_cell_lines_preserve(cell, lines, trim_extra_paragraphs=True)


def _fill_experience_highlights_row(row, item: dict) -> None:
    cells = _unique_cells(row)
    if not cells:
        return
    cell = cells[0]
    highlights = [h.strip() for h in item.get("highlights") or [] if h.strip()]
    lines = _filter_empty_lines(highlights)
    _set_cell_lines_preserve(cell, lines, trim_extra_paragraphs=True)


def _fill_education_row(row, item: dict) -> None:
    cells = _unique_cells(row)
    if not cells:
        return
    left = cells[0]
    right = cells[1] if len(cells) > 1 else None

    institution = (item.get("institution") or "").strip()
    degree = (item.get("degree") or "").strip()
    honors = (item.get("honors") or "").strip()
    honors_line = f"Honores: {honors}" if honors else ""
    left_lines = _filter_empty_lines([institution, degree, honors_line])
    _set_cell_lines_preserve(left, left_lines, trim_extra_paragraphs=True)

    if right is not None:
        location = _join_location(item.get("city"), item.get("country"))
        date_range = _format_date_range(item.get("start"), item.get("end"))
        right_lines = _filter_empty_lines([location, date_range])
        _set_cell_lines_preserve(right, right_lines, trim_extra_paragraphs=True)


def _fill_extra_row(row, items: list[str]) -> None:
    cells = _unique_cells(row)
    if not cells:
        return
    cell = cells[0]
    lines = _filter_empty_lines(items)
    _set_cell_lines_preserve(cell, lines, trim_extra_paragraphs=True)


def _set_contact_row(row, basics: dict) -> None:
    # Construye la fila de contacto con enlaces clickeables
    cells = _unique_cells(row)
    if not cells:
        return
    cell = cells[0]
    template_paragraphs = list(cell.paragraphs)
    template_paragraph = template_paragraphs[0] if template_paragraphs else None
    run_template = None
    for para in template_paragraphs:
        if para.runs:
            run_template = para.runs[0]
            break

    _clear_cell(cell)

    city = basics.get("city", "").strip()
    country = basics.get("country", "").strip()
    linkedin = basics.get("linkedin", "").strip()
    phone = basics.get("phone", "").strip()
    email = basics.get("email", "").strip()
    github = basics.get("github", "").strip()

    location = ", ".join([part for part in [city, country] if part])
    parts: list[ContactPart] = []
    if location:
        parts.append(("text", location, None, 11))
    if linkedin:
        parts.append(("link", linkedin, _normalize_url(linkedin), 10))
    if phone:
        parts.append(("text", phone, None, 11))
    if email:
        parts.append(("link", email, f"mailto:{email}", 11))

    if parts:
        paragraph = cell.add_paragraph()
        if template_paragraph is not None:
            _clone_paragraph_format(template_paragraph, paragraph)
        for idx, (kind, text, url, size_pt) in enumerate(parts):
            if idx > 0:
                _add_run_with_size(paragraph, " · ", run_template, size_pt=11)
            if kind == "link" and url:
                _add_hyperlink(paragraph, text, url, run_template, size_pt=size_pt)
            else:
                _add_run_with_size(paragraph, text, run_template, size_pt=size_pt)

    if github:
        paragraph = cell.add_paragraph()
        if template_paragraph is not None:
            _clone_paragraph_format(template_paragraph, paragraph)
        _add_hyperlink(paragraph, github, _normalize_url(github), run_template, size_pt=10)
        spacer = cell.add_paragraph()
        if template_paragraph is not None:
            _clone_paragraph_format(template_paragraph, spacer)
        _add_run_with_size(spacer, "", run_template, size_pt=11)

    if not parts and not github:
        paragraph = cell.add_paragraph()
        if template_paragraph is not None:
            _clone_paragraph_format(template_paragraph, paragraph)
        _add_run_with_size(paragraph, "", run_template, size_pt=10)


def _normalize_url(value: str) -> str:
    # Asegura que el link tenga protocolo
    raw = value.strip()
    if not raw:
        return raw
    if raw.startswith("http://") or raw.startswith("https://"):
        return raw
    return f"https://{raw}"


def _add_run_with_size(paragraph, text: str, run_template=None, *, size_pt: int | None = None):
    run = paragraph.add_run(text)
    _clone_run_format(run_template, run)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    return run


def _add_hyperlink(
    paragraph,
    text: str,
    url: str,
    run_template=None,
    *,
    size_pt: int | None = None,
) -> None:
    # Inserta un hyperlink con estilo similar al template
    if not url:
        _add_run_with_size(paragraph, text, run_template, size_pt=size_pt)
        return

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    run = Run(new_run, paragraph)
    _clone_run_format(run_template, run)
    run.text = text
    run.underline = True
    run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def _format_date_range(start: str | None, end: str | None) -> str:
    start = _format_date_token(start)
    end = _format_date_token(end)
    if start and end:
        return f"{start}–{end}"
    if start and not end:
        return f"{start}–Actualidad"
    if end and not start:
        return end
    return ""


def _format_date_token(value: str | None) -> str:
    value = (value or "").strip()
    if not value:
        return ""
    match = re.match(r"^(?P<year>\d{4})-(?P<month>\d{2})$", value)
    if match:
        month = MONTHS_REVERSE.get(match.group("month"), match.group("month"))
        return f"{month} {match.group('year')}"
    return value


def _join_location(city: str | None, country: str | None) -> str:
    city = (city or "").strip()
    country = (country or "").strip()
    return ", ".join([part for part in [city, country] if part])


def _set_row_text(row, text: str) -> None:
    cells = _unique_cells(row)
    if not cells:
        return
    lines = _filter_empty_lines([text])
    _set_cell_lines_preserve(cells[0], lines, trim_extra_paragraphs=True)


def _has_experience_content(item: dict) -> bool:
    if not item:
        return False
    fields = [
        item.get("role"),
        item.get("company"),
        item.get("start"),
        item.get("end"),
        item.get("city"),
        item.get("country"),
        item.get("technologies"),
    ]
    if any((str(value).strip() for value in fields if value is not None)):
        return True
    highlights = item.get("highlights") or []
    return any(str(value).strip() for value in highlights)


def _has_education_content(item: dict) -> bool:
    if not item:
        return False
    fields = [
        item.get("degree"),
        item.get("institution"),
        item.get("start"),
        item.get("end"),
        item.get("city"),
        item.get("country"),
        item.get("honors"),
    ]
    return any((str(value).strip() for value in fields if value is not None))


def _set_row_keep_with_next(row, value: bool = True) -> None:
    cells = _unique_cells(row)
    if not cells:
        return
    for cell in cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_with_next = value


def _clear_row_height(row) -> None:
    tr_pr = row._tr.trPr
    if tr_pr is None:
        return
    for child in list(tr_pr):
        if child.tag == qn("w:trHeight"):
            tr_pr.remove(child)


def _apply_section_keep_with_next_gap(table, header_idx: int) -> None:
    gap_idx = header_idx + 1
    if gap_idx < len(table.rows) and _is_blank_row(table.rows[gap_idx]):
        _set_row_keep_with_next(table.rows[gap_idx])
        _clear_row_height(table.rows[gap_idx])


def _set_cell_lines_preserve(cell, lines: list[str], *, trim_extra_paragraphs: bool = False) -> None:
    paragraphs = list(cell.paragraphs)
    template_paragraphs = list(paragraphs)
    if not paragraphs:
        _add_paragraph(cell, lines[0] if lines else "")
        paragraphs = list(cell.paragraphs)
        template_paragraphs = list(paragraphs)

    fallback_run = None
    for para in template_paragraphs:
        if para.runs:
            fallback_run = para.runs[0]
            break

    def pick_template(index: int):
        if not template_paragraphs:
            return None
        if index < len(template_paragraphs):
            return template_paragraphs[index]
        return template_paragraphs[-1]

    for idx, line in enumerate(lines):
        if idx < len(paragraphs):
            template = pick_template(idx)
            run_template = template.runs[0] if template is not None and template.runs else fallback_run
            _set_paragraph_text(paragraphs[idx], line, run_template=run_template)
        else:
            template = pick_template(idx)
            run_template = template.runs[0] if template is not None and template.runs else fallback_run
            _add_paragraph(
                cell,
                line,
                style=template.style if template is not None else None,
                run_template=run_template,
                paragraph_template=template,
            )

    for idx in range(len(lines), len(paragraphs)):
        template = pick_template(idx)
        run_template = template.runs[0] if template is not None and template.runs else fallback_run
        if trim_extra_paragraphs:
            paragraphs[idx]._element.getparent().remove(paragraphs[idx]._element)
            continue
        if _paragraph_has_numbering(paragraphs[idx]):
            paragraphs[idx]._element.getparent().remove(paragraphs[idx]._element)
        else:
            _set_paragraph_text(paragraphs[idx], "", run_template=run_template)


def _set_paragraph_text(paragraph, text: str, run_template=None) -> None:
    if run_template is None and paragraph.runs:
        run_template = paragraph.runs[0]
    _clear_paragraph_content(paragraph)
    run = paragraph.add_run(text)
    _clone_run_format(run_template, run)


def _clear_cell(cell) -> None:
    for paragraph in list(cell.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)


def _clear_paragraph_content(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag == qn("w:pPr"):
            continue
        p.remove(child)


def _paragraph_has_numbering(paragraph) -> bool:
    ppr = paragraph._p.pPr
    return ppr is not None and ppr.numPr is not None


def _add_paragraph(
    cell,
    text: str,
    *,
    bold: bool | None = None,
    style=None,
    run_template=None,
    paragraph_template=None,
) -> None:
    paragraph = cell.add_paragraph()
    if paragraph_template is not None:
        _clone_paragraph_format(paragraph_template, paragraph)
    elif style is not None:
        paragraph.style = style
    run = paragraph.add_run(text)
    _clone_run_format(run_template, run)
    if bold is not None:
        run.bold = bold


def _filter_empty_lines(lines: list[str]) -> list[str]:
    cleaned = []
    for line in lines or []:
        if line is None:
            continue
        trimmed = str(line).strip()
        if trimmed:
            cleaned.append(trimmed)
    return cleaned or [""]


def _clone_run_format(source, target) -> None:
    if source is None:
        return
    target.bold = source.bold
    target.italic = source.italic
    target.underline = source.underline
    if source.font is not None:
        target.font.name = source.font.name
        target.font.size = source.font.size
        if source.font.color is not None and source.font.color.rgb is not None:
            target.font.color.rgb = source.font.color.rgb


def _clone_paragraph_format(source, target) -> None:
    if source is None:
        return
    source_p = source._p
    target_p = target._p
    if source_p.pPr is None:
        return
    if target_p.pPr is not None:
        target_p.remove(target_p.pPr)
    target_p.append(deepcopy(source_p.pPr))


def _set_numbering_level_size(doc: DocxDocumentType, num_id: str, ilvl: str, size_pt: int) -> None:
    numbering = doc.part.numbering_part.element
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    abstract_id = None
    for num in numbering.findall("w:num", ns):
        if num.get(qn("w:numId")) == str(num_id):
            abstract = num.find("w:abstractNumId", ns)
            if abstract is not None:
                abstract_id = abstract.get(qn("w:val"))
            break
    if abstract_id is None:
        return

    abstract_num = None
    for item in numbering.findall("w:abstractNum", ns):
        if item.get(qn("w:abstractNumId")) == str(abstract_id):
            abstract_num = item
            break
    if abstract_num is None:
        return

    level = None
    for lvl in abstract_num.findall("w:lvl", ns):
        if lvl.get(qn("w:ilvl")) == str(ilvl):
            level = lvl
            break
    if level is None:
        return

    rpr = level.find("w:rPr", ns)
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        level.append(rpr)

    size_val = str(int(size_pt * 2))
    for tag in ("w:sz", "w:szCs"):
        el = rpr.find(tag, ns)
        if el is None:
            el = OxmlElement(tag)
            rpr.append(el)
        el.set(qn("w:val"), size_val)


def _apply_font(doc: DocxDocumentType, font_name: str | None) -> None:
    # Aplica la misma fuente a todo el documento
    if not font_name:
        return
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            _set_run_font_name(run, font_name)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _set_run_font_name(run, font_name)


def _set_run_font_name(run, font_name: str) -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _unique_cells(row) -> list:
    seen = set()
    unique = []
    for cell in row.cells:
        key = id(cell._tc)
        if key in seen:
            continue
        seen.add(key)
        unique.append(cell)
    return unique


def _row_text(row) -> str:
    return " ".join(cell.text for cell in row.cells if cell.text)


def _row_is_heading(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    letters = [char for char in stripped if char.isalpha()]
    if letters and all(char.isupper() for char in letters):
        return True
    return False


def _is_blank_row(row) -> bool:
    return not _row_text(row).strip()


def _row_has_borders(row) -> bool:
    xml = row._tr.xml
    return "w:tcBorders" in xml or "w:trBorders" in xml


def _find_blank_row_without_borders(table, start: int, end: int | None) -> int | None:
    limit = end if end is not None else len(table.rows)
    for idx in range(start, min(limit, len(table.rows))):
        row = table.rows[idx]
        if _is_blank_row(row) and not _row_has_borders(row):
            return idx
    return None


def _find_any_blank_row_without_borders(table) -> int | None:
    return _find_blank_row_without_borders(table, 0, None)


def _find_trailing_blank_row_without_borders(table, start: int, end: int | None) -> int | None:
    limit = end if end is not None else len(table.rows)
    last_idx = min(limit, len(table.rows)) - 1
    idx = last_idx
    while idx >= start:
        row = table.rows[idx]
        if not _is_blank_row(row):
            break
        if not _row_has_borders(row):
            return idx
        idx -= 1
    return None


def _collapse_blank_rows(table) -> None:
    idx = 0
    while idx < len(table.rows) - 1:
        row = table.rows[idx]
        next_row = table.rows[idx + 1]
        if (
            _is_blank_row(row)
            and _is_blank_row(next_row)
            and not _row_has_borders(row)
            and not _row_has_borders(next_row)
        ):
            next_row._tr.getparent().remove(next_row._tr)
            continue
        idx += 1


def _find_row_index(table, marker: str) -> int | None:
    target = marker.lower()
    for idx, row in enumerate(table.rows):
        row_text = _row_text(row)
        if target in row_text.lower() and _row_is_heading(row_text):
            return idx
    return None


def _find_row_index_predicate(table, predicate) -> int | None:
    for idx, row in enumerate(table.rows):
        if predicate(_row_text(row).lower()):
            return idx
    return None


def _find_next_non_empty_row(table, start: int, end: int | None) -> int | None:
    limit = end if end is not None else len(table.rows)
    for idx in range(start, min(limit, len(table.rows))):
        if not _is_blank_row(table.rows[idx]):
            return idx
    return None


def _find_first_non_empty_before(table, end: int) -> int | None:
    for idx in range(min(end, len(table.rows))):
        if not _is_blank_row(table.rows[idx]):
            return idx
    return None


def _remove_rows(table, start: int, end: int) -> None:
    for _ in range(max(0, end - start)):
        row = table.rows[start]
        row._tr.getparent().remove(row._tr)


def _insert_row_before(table, row_idx: int, tr_element) -> None:
    rows = list(table._tbl.tr_lst)
    if row_idx >= len(rows):
        table._tbl.append(tr_element)
        return
    rows[row_idx].addprevious(tr_element)
