import io
import unicodedata
from pathlib import Path

from django.test import SimpleTestCase
from docx import Document as DocxDocument

from editor.docx_template import render_from_template


def _normalize(value: str) -> str:
    decomposed = unicodedata.normalize("NFD", value or "")
    without_marks = "".join(ch for ch in decomposed if unicodedata.category(ch) != "Mn")
    return " ".join(without_marks.upper().split())


def _row_text(row) -> str:
    return " ".join(cell.text for cell in row.cells if cell.text).strip()


def _find_heading_index(table, heading: str) -> int | None:
    target = _normalize(heading)
    for idx, row in enumerate(table.rows):
        if target in _normalize(_row_text(row)):
            return idx
    return None


def _find_next_non_empty_row_index(table, start_idx: int) -> int | None:
    for idx in range(start_idx, len(table.rows)):
        if _row_text(table.rows[idx]):
            return idx
    return None


class SkillsPaginationTests(SimpleTestCase):
    def test_skill_categories_only_keep_with_next_from_second_category(self) -> None:
        template_path = Path(__file__).resolve().parents[2] / "templates" / "cv_template.docx"
        structured = {
            "basics": {
                "name": "Test User",
                "description": "Perfil de prueba",
                "email": "",
                "phone": "",
                "linkedin": "",
                "github": "",
                "country": "",
                "city": "",
            },
            "experience": [
                {
                    "role": "Developer",
                    "company": "ACME",
                    "start": "2020-01",
                    "end": "2021-01",
                    "city": "Santiago",
                    "country": "Chile",
                    "technologies": "",
                    "highlights": [],
                }
            ],
            "education": [
                {
                    "degree": "Ingenieria",
                    "institution": "Universidad X",
                    "start": "2010-01",
                    "end": "2014-01",
                    "city": "Santiago",
                    "country": "Chile",
                    "honors": "",
                }
            ],
            "skills": [
                {
                    "category": "Idiomas",
                    "items": "Español nativo, Ingles intermedio",
                },
                {
                    "category": "Herramientas y metodologias",
                    "items": "Git, ESLint, Prettier, Agile",
                }
            ],
            "extra_sections": [],
            "meta": {"core_order": "experience,education,skills"},
        }

        output = render_from_template(structured, template_path)
        doc = DocxDocument(io.BytesIO(output))
        table = doc.tables[0]

        skills_header_idx = _find_heading_index(table, "habilidades")
        self.assertIsNotNone(skills_header_idx)
        assert skills_header_idx is not None

        skills_content_idx = _find_next_non_empty_row_index(table, skills_header_idx + 1)
        self.assertIsNotNone(skills_content_idx)
        assert skills_content_idx is not None

        cell = table.rows[skills_content_idx].cells[0]
        paragraphs = [p for p in cell.paragraphs if p.text and p.text.strip()]
        first_category = None
        target_category = None
        for paragraph in paragraphs:
            normalized = _normalize(paragraph.text)
            if normalized == _normalize("Idiomas"):
                first_category = paragraph
            if normalized == _normalize("Herramientas y metodologias"):
                target_category = paragraph

        self.assertIsNotNone(first_category)
        self.assertIsNotNone(target_category)
        assert first_category is not None and target_category is not None
        self.assertFalse(bool(first_category.paragraph_format.keep_with_next))
        self.assertTrue(bool(target_category.paragraph_format.keep_with_next))
