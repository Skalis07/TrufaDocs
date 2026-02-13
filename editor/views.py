import io
import os
from pathlib import Path
import tempfile
import unicodedata
import zipfile
from datetime import date
from xml.etree import ElementTree


from django.conf import settings
from django.http import HttpResponse
from django.shortcuts import redirect, render
from django.utils.text import slugify
from django.views.decorators.http import require_http_methods
from docx import Document as DocxDocument
from .pdf_parse import parse_pdf_to_structure
from .structure import (
    build_text_from_structure,
    default_structure,
    parse_resume,
    structure_from_post,
)
from .docx_template import render_from_template

# Tipos de archivo permitidos para upload
ALLOWED_EXTENSIONS = {".docx", ".pdf"}
# Fuentes disponibles en la UI para exportar
FONT_CHOICES = [
    "STIX Two Text",
    "Calibri",
    "Arial",
    "Times New Roman",
    "Georgia",
    "Garamond",
    "Montserrat",
    "Poppins",
]

CURRENT_YEAR = date.today().year
YEAR_CHOICES = list(range(CURRENT_YEAR + 5, 1969, -1))

# Listas de paises (minimo viable para el selector)
COUNTRY_CHOICES = [
    "Chile",
    "Venezuela",
]

UI_MESSAGES = {
    "es": {
        "upload_select_file": "Selecciona un archivo .docx o .pdf.",
        "upload_file_too_large": "El archivo supera {max_mb} MB.",
        "upload_unsupported_format": "Formato no soportado. Usa .docx o .pdf.",
        "upload_extract_text_failed": "No se pudo extraer texto del archivo.",
        "export_template_not_found": "No se encontro una plantilla DOCX valida para exportar.",
        "export_docx_template_failed": "No se pudo generar el DOCX desde la plantilla: {detail}",
        "export_docx_template_failed_generic": "No se pudo generar el DOCX desde la plantilla.",
        "export_pdf_failed": "No se pudo generar el PDF.",
        "unknown_error": "error desconocido.",
        "docx_empty": "El archivo DOCX esta vacio.",
        "docx_read_failed": "No se pudo leer el DOCX: {detail}",
        "docx_read_xml_failed": "No se pudo leer el DOCX (XML): {detail}",
        "docx_no_text": "No se encontro texto legible dentro del DOCX.",
        "pdf_read_failed": "No se pudo leer el PDF: {detail}",
        "pdf_no_text": "No se pudo extraer texto del PDF.",
        "docx2pdf_not_installed": "docx2pdf no esta instalado. Ejecuta pip install -r requirements.txt.",
        "docx2pdf_convert_failed_detail": "docx2pdf fallo al convertir: {detail}",
        "docx2pdf_convert_failed_word": "docx2pdf fallo al convertir. Asegura que Microsoft Word este instalado.",
        "docx2pdf_output_missing": "docx2pdf no genero el PDF esperado. Verifica Microsoft Word.",
    },
    "en": {
        "upload_select_file": "Please select a .docx or .pdf file.",
        "upload_file_too_large": "The file exceeds {max_mb} MB.",
        "upload_unsupported_format": "Unsupported format. Use .docx or .pdf.",
        "upload_extract_text_failed": "Could not extract text from the file.",
        "export_template_not_found": "No valid DOCX template was found for export.",
        "export_docx_template_failed": "Could not generate the DOCX from the template: {detail}",
        "export_docx_template_failed_generic": "Could not generate the DOCX from the template.",
        "export_pdf_failed": "Could not generate the PDF.",
        "unknown_error": "unknown error.",
        "docx_empty": "The DOCX file is empty.",
        "docx_read_failed": "Could not read the DOCX: {detail}",
        "docx_read_xml_failed": "Could not read the DOCX (XML): {detail}",
        "docx_no_text": "No readable text was found inside the DOCX.",
        "pdf_read_failed": "Could not read the PDF: {detail}",
        "pdf_no_text": "Could not extract text from the PDF.",
        "docx2pdf_not_installed": "docx2pdf is not installed. Run pip install -r requirements.txt.",
        "docx2pdf_convert_failed_detail": "docx2pdf failed to convert: {detail}",
        "docx2pdf_convert_failed_word": "docx2pdf failed to convert. Make sure Microsoft Word is installed.",
        "docx2pdf_output_missing": "docx2pdf did not generate the expected PDF. Check Microsoft Word.",
    },
}


def _ui_lang(request) -> str:
    raw = ""
    if getattr(request, "method", "") == "POST":
        raw = request.POST.get("ui_lang", "")
    if not raw:
        raw = request.GET.get("ui_lang", "")
    return "en" if str(raw).strip().lower() == "en" else "es"


def _msg(request, key: str, **kwargs) -> str:
    lang = _ui_lang(request)
    template = UI_MESSAGES.get(lang, UI_MESSAGES["es"]).get(key) or UI_MESSAGES["es"].get(key) or key
    if not kwargs:
        return template
    try:
        return template.format(**kwargs)
    except Exception:
        return template


def _translate_backend_error(request, message: str | None) -> str | None:
    if not message:
        return message
    text = str(message).strip()
    if not text:
        return text

    exact_map = {
        "El archivo DOCX esta vacio.": "docx_empty",
        "No se encontro texto legible dentro del DOCX.": "docx_no_text",
        "No se pudo extraer texto del PDF.": "pdf_no_text",
        "docx2pdf no esta instalado. Ejecuta pip install -r requirements.txt.": "docx2pdf_not_installed",
        "docx2pdf fallo al convertir. Asegura que Microsoft Word este instalado.": "docx2pdf_convert_failed_word",
        "docx2pdf no genero el PDF esperado. Verifica Microsoft Word.": "docx2pdf_output_missing",
    }
    mapped_key = exact_map.get(text)
    if mapped_key:
        return _msg(request, mapped_key)

    prefix_map = [
        ("No se pudo leer el DOCX: ", "docx_read_failed"),
        ("No se pudo leer el DOCX (XML): ", "docx_read_xml_failed"),
        ("No se pudo leer el PDF: ", "pdf_read_failed"),
        ("docx2pdf fallo al convertir: ", "docx2pdf_convert_failed_detail"),
    ]
    for prefix, key in prefix_map:
        if text.startswith(prefix):
            detail = text[len(prefix) :].strip() or _msg(request, "unknown_error")
            return _msg(request, key, detail=detail)

    return text

def _extract_structured_countries(structured: dict | None) -> list[str]:
    # Recoge paises desde la estructura para que se muestren en el selector.
    countries: list[str] = []

    if not isinstance(structured, dict):
        structured = {}

    def add(country: str | None) -> None:
        if not country:
            return
        country_clean = country.strip()
        if country_clean:
            countries.append(country_clean)

    basics = structured.get("basics") or {}
    add(basics.get("country"))

    for item in structured.get("experience") or []:
        add(item.get("country"))
    for item in structured.get("education") or []:
        add(item.get("country"))
    for extra in structured.get("extra_sections") or []:
        for entry in extra.get("entries") or []:
            add(entry.get("country"))

    return countries


def _merge_country_choices(base: list[str], extra: list[str]) -> list[str]:
    merged = list(base)
    seen = {c.strip().lower() for c in merged if c}
    for country in extra:
        key = country.strip().lower()
        if key and key not in seen:
            merged.append(country.strip())
            seen.add(key)
    return merged


def index(request):
    # Estado inicial: formulario con estructura vacia
    structured = default_structure()
    return _render_text_editor(request, structured, filename="documento")
# --------------------
# Modo texto (opcional)
# --------------------



@require_http_methods(["GET", "POST"])
def text_upload(request):
    if request.method == "GET":
        return redirect("index")

    uploaded = request.FILES.get("file")
    if not uploaded:
        return _text_error(request, _msg(request, "upload_select_file"))

    max_mb = _max_upload_mb()
    if uploaded.size > max_mb * 1024 * 1024:
        return _text_error(request, _msg(request, "upload_file_too_large", max_mb=max_mb))

    if not _is_allowed_extension(uploaded.name):
        return _text_error(request, _msg(request, "upload_unsupported_format"))

    ext = _extension(uploaded.name)
    if ext == ".docx":
        text, error = _extract_docx_text(uploaded)
        if error:
            return _text_error(request, _translate_backend_error(request, error) or error)
        if not text.strip():
            return _text_error(request, _msg(request, "upload_extract_text_failed"))
        structured = parse_resume(text)
    else:
        structured, error = parse_pdf_to_structure(uploaded)
        if error:
            return _text_error(request, _translate_backend_error(request, error) or error)

    filename = _safe_filename(uploaded.name)
    return _render_text_editor(request, structured, filename=filename)


@require_http_methods(["POST"])
def export_docx(request):
    raw_text = request.POST.get("text", "")
    use_structured = request.POST.get("use_structured") == "1"
    if use_structured:
        # Exporta usando la estructura (plantilla DOCX)
        structured = structure_from_post(request.POST)
        template_path = _template_path()
        font_choice = _selected_font(request)
        if not template_path:
            return _render_text_editor(
                request,
                structured,
                filename=_safe_filename(request.POST.get("filename", "documento")),
                error=_msg(request, "export_template_not_found"),
            )
        try:
            rendered = render_from_template(
                structured,
                template_path,
                font_name=font_choice,
                ui_lang=_ui_lang(request),
            )
        except Exception as exc:
            detail = str(exc).strip()
            if len(detail) > 400:
                detail = detail[:400].rstrip() + "..."
            return _render_text_editor(
                request,
                structured,
                filename=_safe_filename(request.POST.get("filename", "documento")),
                error=_msg(
                    request,
                    "export_docx_template_failed",
                    detail=detail or _msg(request, "unknown_error"),
                ),
            )
        filename = _safe_filename(request.POST.get("filename", "documento"))
        response = HttpResponse(
            rendered,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}.docx"'
        return response
    else:
        # Exporta el texto libre (columna derecha)
        text = raw_text
    filename = _safe_filename(request.POST.get("filename", "documento"))

    docx_bytes = _build_docx_bytes(text)

    response = HttpResponse(
        docx_bytes,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}.docx"'
    return response


@require_http_methods(["POST"])
def export_pdf(request):
    raw_text = request.POST.get("text", "")
    use_structured = request.POST.get("use_structured") == "1"
    filename = _safe_filename(request.POST.get("filename", "documento"))

    if use_structured:
        # PDF desde estructura -> DOCX -> PDF
        structured = structure_from_post(request.POST)
        template_path = _template_path()
        font_choice = _selected_font(request)
        if not template_path:
            return _render_text_editor(
                request,
                structured,
                filename=filename,
                error=_msg(request, "export_template_not_found"),
            )
        try:
            rendered_docx = render_from_template(
                structured,
                template_path,
                font_name=font_choice,
                ui_lang=_ui_lang(request),
            )
        except Exception:
            rendered_docx = None
        if not rendered_docx:
            return _render_text_editor(
                request,
                structured,
                filename=filename,
                error=_msg(request, "export_docx_template_failed_generic"),
            )
        docx_bytes = rendered_docx
    else:
        # PDF desde texto libre
        structured = default_structure()
        docx_bytes = _build_docx_bytes(raw_text)

    pdf_bytes, error = _convert_docx_bytes_to_pdf(docx_bytes)
    if pdf_bytes:
        response = HttpResponse(pdf_bytes, content_type="application/pdf")
        response["Content-Disposition"] = f'attachment; filename="{filename}.pdf"'
        return response
    return _render_text_editor(
        request,
        structured,
        filename=filename,
        error=_translate_backend_error(request, error) or _msg(request, "export_pdf_failed"),
    )


# --------------------
# Helpers
# --------------------


def _render_text_editor(request, structured, filename="documento", error: str | None = None):
    # Render principal con datos de la UI
    font_choice = _selected_font(request)
    extra_countries = _extract_structured_countries(structured)
    country_choices = _merge_country_choices(COUNTRY_CHOICES, extra_countries)
    return render(
        request,
        "editor/editor.html",
        {
            "text": build_text_from_structure(structured),
            "filename": filename,
            "structured": structured,
            "error": error,
            "ui_lang": _ui_lang(request),
            "font_choices": FONT_CHOICES,
            "selected_font": font_choice,
            "country_choices": country_choices,
            "year_choices": YEAR_CHOICES,
        },
    )


def _text_error(request, message: str):
    structured = default_structure()
    return _render_text_editor(request, structured, filename="documento", error=message)


def _is_allowed_extension(filename: str) -> bool:
    ext = _extension(filename)
    return ext in ALLOWED_EXTENSIONS


def _max_upload_mb() -> int:
    return int(getattr(settings, "MAX_UPLOAD_MB", 25))


def _extension(name: str) -> str:
    return os.path.splitext(name)[1].lower()


def _safe_filename(name: str) -> str:
    base = os.path.splitext(os.path.basename(name))[0] or "documento"
    safe = slugify(base)
    return safe[:80] or "documento"


def _build_docx_bytes(text: str) -> bytes:
    # DOCX basico para exportaciones sin plantilla
    buffer = io.BytesIO()
    doc = DocxDocument()
    for line in text.splitlines():
        doc.add_paragraph(line)
    if not text.strip():
        doc.add_paragraph("")
    doc.save(buffer)
    return buffer.getvalue()


def _extract_docx_text(file_obj) -> tuple[str, str | None]:
    # Extrae texto de un .docx (usa python-docx)
    raw = file_obj.read()
    if not raw:
        return "", "El archivo DOCX esta vacio."

    lines: list[str] = []
    try:
        doc = DocxDocument(io.BytesIO(raw))
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_cells: list[list[str]] = []
                for cell in row.cells:
                    cell_lines = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
                    if cell_lines:
                        row_cells.append(cell_lines)

                if not row_cells:
                    continue

                seen_cells: set[str] = set()
                for cell_lines in row_cells:
                    key = _normalize_key(" ".join(cell_lines))
                    if not key or key in seen_cells:
                        continue
                    seen_cells.add(key)
                    lines.extend(cell_lines)
    except Exception as exc:
        return "", f"No se pudo leer el DOCX: {str(exc).strip() or 'error desconocido.'}"

    if lines:
        return "\n".join(lines), None

    # Fallback: parse raw XML (text boxes, shapes, etc.)
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as zf:
            candidates = [
                name
                for name in zf.namelist()
                if name == "word/document.xml" or name.startswith("word/header")
            ]
            for name in candidates:
                xml_bytes = zf.read(name)
                try:
                    root = ElementTree.fromstring(xml_bytes)
                except ElementTree.ParseError:
                    continue
                for para in root.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"):
                    texts = [
                        node.text
                        for node in para.findall(
                            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
                        )
                        if node.text
                    ]
                    if texts:
                        line = "".join(texts).strip()
                        if line:
                            lines.append(line)
    except Exception as exc:
        return "", f"No se pudo leer el DOCX (XML): {str(exc).strip() or 'error desconocido.'}"

    if lines:
        return "\n".join(lines), None
    return "", "No se encontro texto legible dentro del DOCX."


def _template_path() -> Path | None:
    # Ruta de plantilla definida por .env o default
    env_path = getattr(settings, "CV_TEMPLATE_PATH", "").strip()
    if env_path:
        candidate = Path(env_path)
        if not candidate.is_absolute():
            candidate = Path(settings.BASE_DIR) / candidate
        if candidate.is_file():
            return candidate

    candidate = Path(settings.BASE_DIR) / "templates" / "cv_template.docx"
    if candidate.is_file():
        return candidate
    return None


def _selected_font(request) -> str:
    # Solo valida fuentes permitidas
    if request.method != "POST":
        return ""
    choice = request.POST.get("doc_font", "").strip()
    if not choice:
        return ""
    if choice in FONT_CHOICES:
        return choice
    return ""


def _normalize_key(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", text)
    return "".join(char for char in normalized if not unicodedata.combining(char)).lower().strip()


def _convert_docx_bytes_to_pdf(docx_bytes: bytes) -> tuple[bytes | None, str | None]:
    # Convierte DOCX a PDF usando docx2pdf (requiere Word instalado)
    try:
        from docx2pdf import convert as docx2pdf_convert
    except Exception:
        return None, "docx2pdf no esta instalado. Ejecuta pip install -r requirements.txt."

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        input_path = tmp_path / "documento.docx"
        input_path.write_bytes(docx_bytes)
        output_path = tmp_path / "documento.pdf"

        pythoncom = None
        try:
            import pythoncom  # type: ignore
        except Exception:
            pythoncom = None

        co_initialized = False
        try:
            if pythoncom is not None:
                pythoncom.CoInitialize()
                co_initialized = True
            docx2pdf_convert(str(input_path), str(output_path))
        except Exception as exc:
            detail = str(exc).strip()
            if detail:
                if len(detail) > 400:
                    detail = detail[:400].rstrip() + "..."
                return None, f"docx2pdf fallo al convertir: {detail}"
            return None, "docx2pdf fallo al convertir. Asegura que Microsoft Word este instalado."
        finally:
            if pythoncom is not None and co_initialized:
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass

        if not output_path.exists():
            alt_path = input_path.with_suffix(".pdf")
            if alt_path.exists():
                output_path = alt_path
            else:
                return None, "docx2pdf no genero el PDF esperado. Verifica Microsoft Word."

        return output_path.read_bytes(), None
